import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
from io import BytesIO
import math

load_dotenv()

st.set_page_config(
    page_title="Language Translation Agent",
    page_icon="🗣️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    [data-testid="stExpanderDetails"] {
        padding: 0rem;
      }
</style>
""", unsafe_allow_html=True)

CHUNK_SIZE = 2

DEFAULT_PROMPT_INSTRUCTIONS = """You are a highly skilled translation expert. Your task is to translate the provided English text into the specified target language.
- Your output must ONLY be the translated text itself.
- Do not include any introductory phrases like "Here is the translation:" or any other conversational filler.
- Preserve the original formatting (like paragraphs and line breaks) as much as possible."""

@st.cache_resource
def load_model(model_name):
    """Loads the specified Gemini model and caches it."""
    return genai.GenerativeModel(model_name=model_name)

@st.cache_data
def translate_text(_model, text, target_language, instructions):
    """Translates text and caches the result."""
    if not text or not target_language:
        return ""
    try:
        final_prompt = f"""
        {instructions}

        Translate the following English text to **{target_language}**:
        ```
        {text}
        ```
        """
        response = _model.generate_content(final_prompt)
        print("API Call Made")
        return response.text.strip()
    except Exception as e:
        st.error(f"An error occurred during translation: {str(e)}")
        return None

def extract_text_from_pdf(file):
    try:
        pdf_reader = PdfReader(file)
        return [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return []

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return [para.text for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return []

def extract_text_from_txt(file):
    try:
        content = file.getvalue().decode("utf-8")
        lines = content.splitlines()
        return ["\n".join(lines[i:i + 40]) for i in range(0, len(lines), 40)]
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return []

def create_pdf_from_text(text):
    pdf = FPDF()
    pdf.add_page()
    try:
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.set_font("DejaVu", size=12)
    except RuntimeError:
        st.warning("DejaVu font not found. Falling back to default font.")
        pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    return bytes(pdf.output(dest='S'))

def create_docx_from_text(text):
    doc = Document()
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def init_session_state():
    """Initializes session state variables if they don't exist."""
    state_defaults = {
        "api_key_configured": False,
        "text_translation_result": "",
        "doc_translation_result": "",
        "target_language": "Hindi",
        "doc_info": {"name": None, "type": None},
        "selected_model": "models/gemini-1.5-flash"
    }
    for key, value in state_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

api_key = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY"))
if not api_key:
    st.error("Gemini API Key is not set. Please add it to secrets or a .env file.")
    st.stop()

try:
    genai.configure(api_key=api_key)
    st.session_state.api_key_configured = True
except Exception as e:
    st.error(f"Invalid API Key or configuration error: {e}")
    st.session_state.api_key_configured = False
    st.stop()

st.title("🗣️ Language Translation Agent")
st.markdown("**Translate Text or Documents using the power of Gemini.**")

st.markdown("Language and Model Selection")
col_lang, col_model, col_empty = st.columns([1, 1, 2])

languages = ["Abkhaz", "Acehnese", "Acholi", "Afar", "Afrikaans", "Albanian", "Alur", "Amharic", "Arabic", "Armenian", "Assamese", "Avar", "Awadhi", "Aymara","Azerbaijani", "Balinese", "Baluchi","Bambara", "Baoulé", "Bashkir","Basque","Batak Karo","Batak Simalungun","Batak Toba","Belarusian","Bemba","Bengali","Betawi","Bhojpuri","Bikol","Bosnian","Breton","Bulgarian","Buryat", "Cantonese","Catalan","Cebuano","Chamorro","Chechen","Chichewa","Chinese (Simplified)","Chinese (Traditional)","Chuukese","Chuvash","Corsican","Crimean Tatar (Cyrillic)","Crimean Tatar (Latin)","Croatian","Czech","Danish","Dari","Dhivehi","Dinka","Dogri","Dombe","Dutch","Dyula","Dzongkha","English","Esperanto","Estonian","Ewe","Faroese","Fijian","Filipino","Finnish","Fon","French","French (Canada)","Frisian","Friulian","Fulani", "Ga","Galician","Georgian","German","Greek","Guarani","Gujrati","Haitian Crele","Hakha Chin","Hausa","Hawaiian","Hebrew","Hiligaynon","Hindi","Hmong","Hungarian","Hunsrik","Iban","Icelandic","Igbo","Ilocano","Indonesian","Inuktut (Latin)","Inuktut (Syllabics)","Irish","Italian","Jamaican Patois","Japanese","Javanese","Jingpo","Kalaallisut","Kannada","Kanuri","Kapampangan","Kazakh","Khasi","Khmer","Kiga","Kikongo","Kinyarwanda","Kituba","Kokborok","Komi","Konkani","Korean","Krio","Kurdish (Kurmanji)","Kurdish (Sorani)","Kyrgyz",
"Lao","Latgalian","Latin","Latvian","Ligurian","Limburgish","Lingala","Lithuanian","Lombard","Luganda","Luo","Luxembourgish", "Macedonian","Madurese","Maithili","Makassar","Malagasy","Malay","Malay (Jawi)","Malayalam","Maltese","Mam","Manx","Maroi","Marathi","Marshallese","Marwadi","Mauitian Creole","Meadow Mari","Meiteilon (Manipuri)","Minang","Mizo","Mongolian","Myanmar (Burmese)","Nahuatl (Eastern Huasteca)","Ndau","Ndebele (South)","Nepalbhasa (Newari)","Nepali","NKo","Norwegian","Nuer","Occitan","Odia (Oriya)","Oromo","Ossetian","Pangasinan","Papiamento","Pashto","Persian","Polish","Portuguese (Brazil)","Portuguese (Portugal)","Punjabi (Gurmukhi)","Punjabi (Shahmukhi)","Quechua","Qʼeqchiʼ","Romani","Romanian","Rundi","Russian","Sami (North)","Samoan","Sango","Sanskrit","Santali (Latin)","Santali (Ol Chiki)","Scots Gaelic","Sepedi","Serbian","Sesotho","Seychellois Creole","Shan","Shona","Sicilian","Silesian","Sindhi","Sinhala","Slovak","Slovenian","Somali","Spanish","Sundanese", "Susu","Swahili","Swati","Swedish","Tahitian","Tazik","Tamazight","Tamazight (Tifinagh)","Tamil","Tatr","Telugu","Tetum","Thai","Tibetan","Tigrinya","Tiv","Tok Pisin","Tongan","Tshiluba","Tsonga","Tswana","Tulu","Tumbuka","Turkish","Turkmen","Tuvan","Twi","Udmurt","Ukrainian","Urdu","Uyghur","Uzbek","Venda","Venetian","Vietnamese","Waray","Welsh","Wolof","Xhosa","Yakut","Yiddish","Yoruba","Yucatec Maya","Zapotec","Zulu"]
languages.sort()

try:
    default_lang_index = languages.index(st.session_state.target_language)
except ValueError:
    languages.append(st.session_state.target_language)
    languages.sort()
    default_lang_index = languages.index(st.session_state.target_language)

with col_lang:
    st.session_state.target_language = st.selectbox(
        "Select Target Language",
        languages,
        index=default_lang_index,
        label_visibility="collapsed"
    )

valid_model_names = ["models/gemini-1.5-pro", "models/gemini-1.5-flash-8b","gemini-1.5-flash", "models/gemini-2.0-flash-live-001", "models/gemini-2.0-flash-lite", 
"models/gemini-2.0-flash-preview-image-generation", "models/gemini-2.0-flash", "models/gemini-2.5-pro-preview-tts", "models/gemini-2.5-flash-preview-tts", "models/gemini-2.5-flash-image-preview", 
"models/" "models/gemini-live-2.5-flash-preview", "models/gemini-2.5-flash-lite", "models/gemini-2.5-flash", "models/gemini-2.5-pro"]

with col_model:
    try:
        default_model_index = valid_model_names.index(st.session_state.selected_model)
    except ValueError:
        default_model_index = 0

    st.session_state.selected_model = st.selectbox(
        "Select Model",
        options=valid_model_names,
        index=default_model_index,
        format_func=lambda name: name.replace("models/", ""),
        label_visibility="collapsed"
    )

model = load_model(st.session_state.selected_model)

with st.expander("Advanced Options: Customize AI Instructions"):
    custom_instructions = st.text_area(
        "Enter custom instructions:",
        placeholder="e.g., Translate the following text into a formal, business-appropriate tone.",
        height=150,
        label_visibility="collapsed"
    )

st.divider()

st.subheader("Translate Text")
col1a, col1b = st.columns([1, 1], gap="large")

with col1a:
    input_text = st.text_area("Enter English text to translate:", height=300, key="text_input")
    if st.button("Translate Text", use_container_width=True, type="primary", disabled=not input_text):
        instructions = custom_instructions.strip() or DEFAULT_PROMPT_INSTRUCTIONS
        spinner_model_name = st.session_state.selected_model.replace('models/', '')
        with st.spinner(f"Translating using {spinner_model_name}..."):
            translated_output = translate_text(model, input_text, st.session_state.target_language, instructions)
            if translated_output is not None:
                st.session_state.text_translation_result = translated_output
                st.session_state.doc_translation_result = ""
                st.session_state.doc_info = {"name": None, "type": None}

with col1b:
    st.text_area(
        "Text Translation Result",
        value=st.session_state.text_translation_result or "Translation will appear here...",
        height=300,
        disabled=True,
        key="text_output"
    )

st.divider()

st.subheader("Translate Document")
col2a, col2b = st.columns([1, 1], gap="large")

with col2a:
    uploaded_file = st.file_uploader("Upload a document", type=["pdf", "docx", "txt"], key="file_uploader")
    if st.button("Translate Document", use_container_width=True, type="primary", disabled=uploaded_file is None):
        if uploaded_file is not None:
            instructions = custom_instructions.strip() or DEFAULT_PROMPT_INSTRUCTIONS
            st.session_state.doc_info['name'] = uploaded_file.name
            file_ext = os.path.splitext(uploaded_file.name)[-1].lower().strip('.')
            st.session_state.doc_info['type'] = file_ext

            source_chunks = []
            if file_ext == "pdf": source_chunks = extract_text_from_pdf(uploaded_file)
            elif file_ext == "docx": source_chunks = extract_text_from_docx(uploaded_file)
            elif file_ext == "txt": source_chunks = extract_text_from_txt(uploaded_file)

            if source_chunks:
                translated_chunks = []
                total_chunks = math.ceil(len(source_chunks) / CHUNK_SIZE)
                progress_bar = st.progress(0, text=f"Translating chunk 1 of {total_chunks}...")

                for i in range(0, len(source_chunks), CHUNK_SIZE):
                    chunk_group = "\n".join(source_chunks[i:i + CHUNK_SIZE])
                    current_chunk_number = (i // CHUNK_SIZE) + 1
                    progress_text = f"Translating chunk {current_chunk_number} of {total_chunks}..."
                    progress_bar.progress(current_chunk_number / total_chunks, text=progress_text)

                    spinner_model_name = st.session_state.selected_model.replace('models/', '')
                    with st.spinner(f"Translating chunk {current_chunk_number}/{total_chunks} using {spinner_model_name}..."):
                        translated_output = translate_text(model, chunk_group, st.session_state.target_language, instructions)
                    if translated_output is not None:
                        translated_chunks.append(translated_output)
                    else:
                        st.error(f"Failed to translate chunk {current_chunk_number}. Skipping.")

                st.session_state.doc_translation_result = "\n\n".join(translated_chunks)
                st.session_state.text_translation_result = ""
                progress_bar.progress(1.0, text="Translation complete!")
            else:
                st.warning("Could not extract text from the document. The file might be empty or corrupted.")

with col2b:
    st.text_area(
        "Document Translation Result",
        value=st.session_state.doc_translation_result or "Document translation will appear here...",
        height=300,
        disabled=True,
        key="doc_output"
    )
    
    if st.session_state.doc_translation_result:
        original_name = os.path.splitext(st.session_state.doc_info.get('name', 'document'))[0]
        file_name = f"translated_{original_name}.txt"
        
        download_data = st.session_state.doc_translation_result.encode('utf-8')

        st.download_button(
            label="Download Translated .txt",
            data=download_data,
            file_name=file_name,
            mime="text/plain",
            use_container_width=True
        )