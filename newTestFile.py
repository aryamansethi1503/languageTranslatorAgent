import streamlit as st
import google.generativeai as genai
from google.cloud import aiplatform
from dotenv import load_dotenv
import os
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
from io import BytesIO
import math
import psycopg2
import psycopg2.pool
import atexit

load_dotenv()

@st.cache_resource
def get_db_pool():
    """
    Creates and returns a thread-safe connection pool to the AlloyDB database.
    This is configured to connect via the AlloyDB Auth Proxy (running as a sidecar).
    """
    try:
        db_user = os.environ["DB_USER"]
        db_pass = os.environ["DB_PASS"]
        db_name = os.environ["DB_NAME"]
        
        db_host = os.environ.get("DB_HOST", "127.0.0.1") 
        db_port = os.environ.get("DB_PORT", "5432")

        print(f"Creating database connection pool for {db_host}:{db_port}...")

        pool = psycopg2.pool.SimpleConnectionPool(
            minconn=1,
            maxconn=10,
            host=db_host, 
            port=db_port,
            dbname=db_name,
            user=db_user,
            password=db_pass
        )
        
        print("Database connection pool created successfully.")
        
        atexit.register(pool.closeall)
        return pool
    except Exception as e:
        print(f"Error creating database connection pool: {e}")
        st.error(f"FATAL: Could not connect to the database. {e}")
        st.stop()

def db_setup(pool):
    """
    Connects to the PostgreSQL-compatible database (AlloyDB) via the pool
    and creates the 'corrections' table if it doesn't already exist.
    """
    print("Checking database schema...")
    conn = None
    try:
        conn = pool.getconn()
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS corrections (
                id SERIAL PRIMARY KEY,
                source_language TEXT DEFAULT 'English',
                target_language TEXT,
                source_text TEXT,
                original_translation TEXT,
                corrected_translation TEXT,
                timestamp TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(source_language, target_language, source_text, corrected_translation)
            )
        ''')
        conn.commit()
        cursor.close()
        print("Database schema check complete.")
    except Exception as e:
        print(f"Error during DB setup: {e}")
        st.error(f"Error during DB setup: {e}")
    finally:
        if conn:
            pool.putconn(conn)

def db_save_correction(pool, target_language, source_text, original_translation, corrected_translation):
    """
    Saves a new translation correction to the AlloyDB database using the pool.
    """
    conn = None
    try:
        conn = pool.getconn()
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT INTO corrections 
            (source_language, target_language, source_text, original_translation, corrected_translation) 
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (source_language, target_language, source_text, corrected_translation) DO NOTHING
        """, ('English', target_language, source_text, original_translation, corrected_translation))
        conn.commit()
        cursor.close()
        
        print("Correction saved to AlloyDB.")
        st.toast("Correction saved!", icon="✅")
    except Exception as e:
        print(f"Error saving correction: {e}")
        st.error(f"Error saving correction: {e}")
    finally:
        if conn:
            pool.putconn(conn)

def db_get_corrections(pool, target_language, limit=3):
    """
    Retrieves the most recent 'limit' corrections from the AlloyDB database using the pool.
    """
    rows = []
    conn = None
    try:
        conn = pool.getconn()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT source_text, corrected_translation 
            FROM corrections 
            WHERE target_language = %s
            ORDER BY timestamp DESC
            LIMIT %s
        """, (target_language, limit))
        rows = cursor.fetchall()
        cursor.close()
    except Exception as e:
        print(f"Error fetching corrections: {e}")
    finally:
        if conn:
            pool.putconn(conn)
    return rows

st.set_page_config(
    page_title="Language Translation Agent",
    page_icon="🗣️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    /* Remove padding from expander details */
    [data-testid="stExpanderDetails"] { padding: 0rem; }
    
    /* Import Google Sans font */
    @import url('https://fonts.googleapis.com/css2?family=Google+Sans:wght@400;500;700&display=swap');

    /* --- Global Styles --- */
    .stApp {
        font-family: 'Google Sans', sans-serif;
    }

    /* --- Text Area Styles --- */
    [data-testid="stTextArea"] textarea {
        font-size: 14px;
        color: #111; /* Dark text for readability on white */
        background-color: white;
    }
    
    /* Style for the disabled output box to match the input */
    [data-testid="stTextArea"] textarea:disabled {
        font-size: 14px;
        color: #333; /* Slightly lighter text for disabled */
        background-color: #FAFAFA; /* Slightly off-white for disabled */
        border: 1px solid #E0E0E0;
    }
</style>
""", unsafe_allow_html=True)

CHUNK_SIZE = 2
DEFAULT_PROMPT_INSTRUCTIONS = """You are a highly skilled translation expert. Your task is to translate the provided English text into the specified target language.
- Your output must ONLY be the translated text itself.
- Do not include any introductory phrases like "Here is the translation:" or any other conversational filler.
- Preserve the original formatting (like paragraphs and line breaks) as much as possible."""

def init_session_state():
    state_defaults = {
        "gemini_api_configured": False,
        "custom_model_configured": False,
        "text_translation_result": "",
        "doc_translation_result": "",
        "target_language": "Hindi",
        "doc_info": {"name": None, "type": None},
        "selected_model": "models/gemini-2.5-flash",
        "prompt_instructions": DEFAULT_PROMPT_INSTRUCTIONS,
        "editing_translation": False,
        "current_source_text": "",
        "text_output_edit": ""
    }
    for key, value in state_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

db_pool = get_db_pool()

db_setup(db_pool)

init_session_state()

try:
    gemini_api_key = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY"))
    if gemini_api_key:
        genai.configure(api_key=gemini_api_key)
        st.session_state.gemini_api_configured = True
except Exception as e:
    st.warning(f"Could not configure Gemini API: {e}")
    
if "custom_model" in st.secrets:
    try:
        aiplatform.init(
            project=st.secrets["custom_model"]["project_id"],
            location=st.secrets["custom_model"]["location"],
        )
        st.session_state.custom_model_configured = True
    except Exception as e:
        st.warning(f"Custom model credentials found, but failed to initialize Vertex AI SDK: {e}")

if not st.session_state.gemini_api_configured and "custom_model" not in st.secrets:
    st.error("FATAL: No models are configured. Please add credentials for Gemini or your custom model in secrets.toml.")
    st.stop()

@st.cache_resource
def load_gemini_model(model_name):
    """Loads the specified Gemini model and caches it."""
    return genai.GenerativeModel(model_name=model_name)

def build_examples_prompt(target_language):
    """Helper function to fetch corrections and build the prompt block."""
    corrections = db_get_corrections(db_pool, target_language, limit=3)
    examples_prompt_block = ""
    if corrections:
        examples_prompt_block = "Here are some examples of past corrections. Use them as a guide to improve accuracy and follow the user's preferred style:\n"
        for i, (source, corrected) in enumerate(corrections):
            examples_prompt_block += f"\n--- Example {i+1} ---\n"
            examples_prompt_block += f"English Text: \"{source}\"\n"
            examples_prompt_block += f"Correct {target_language} Translation: \"{corrected}\"\n"
        examples_prompt_block += "-------------------------\n\n"
    return examples_prompt_block

@st.cache_data(show_spinner=False)
def translate_with_custom_model(_text, target_language, instructions):
    """Translates text using the custom Vertex AI model."""
    if not st.session_state.custom_model_configured:
        st.error("Custom model selected, but its configuration in secrets.toml is missing or invalid.")
        return None
    try:
        endpoint_id = st.secrets["custom_model"]["endpoint_id"]
        project_id = st.secrets["custom_model"]["project_id"]
        location = st.secrets["custom_model"]["location"]
        
        gemma_endpoint = aiplatform.Endpoint(
            endpoint_name=f"projects/{project_id}/locations/{location}/endpoints/{endpoint_id}"
        )

        examples_prompt_block = build_examples_prompt(target_language)
        
        final_prompt = f"""
        {instructions}
        {examples_prompt_block}
        Translate the following English text to **{target_language}**:
        ```
        {_text}
        ```
        """
        instances = [{"prompt": final_prompt}]
        response = gemma_endpoint.predict(instances=instances)
        
        if response.predictions and len(response.predictions) > 0:
            first_prediction = response.predictions[0]
            translated_text = None
            if 'translation_output' in first_prediction:
                translated_text = first_prediction['translation_output']
            elif 'content' in first_prediction:
                 translated_text = first_prediction['content']
            
            if translated_text:
                print("Vertex AI API Call Made")
                return translated_text.strip()
            else:
                st.error("Response format from custom model not recognized. Could not find 'translation_output' or 'content' key.")
                st.json(response.predictions)
                return None
        else:
            st.error("The custom model responded, but the prediction format was not recognized.")
            st.json(response)
            return None
            
    except Exception as e:
        st.error(f"A critical error occurred while contacting the Vertex AI endpoint: {e}")
        return None

@st.cache_data(show_spinner=False)
def translate_text(model_name, text, target_language, instructions):
    """
    Checks the model name and calls the appropriate translation service,
    injecting corrections from the database.
    """
    if not text or not target_language:
        return ""

    if "custom/" in model_name:
        return translate_with_custom_model(text, target_language, instructions)
    else:
        if not st.session_state.gemini_api_configured:
            st.error("Gemini model selected, but the API key is missing or invalid.")
            return None
        try:
            model = load_gemini_model(model_name)
            examples_prompt_block = build_examples_prompt(target_language)
            
            final_prompt = f"""
            {instructions}
            {examples_prompt_block}
            Translate the following English text to **{target_language}**:
            ```
            {text}
            ```
            """
            response = model.generate_content(final_prompt)
            print("Gemini API Call Made")
            return response.text.strip()
        except Exception as e:
            st.error(f"An error occurred during Gemini translation: {str(e)}")
            return None

def extract_text_from_pdf(file):
    try:
        pdf_reader = PdfReader(file)
        return [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return []

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return [para.text for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return []

def extract_text_from_txt(file):
    try:
        content = file.getvalue().decode("utf-8")
        lines = content.splitlines()
        return ["\n".join(lines[i:i + 40]) for i in range(0, len(lines), 40)]
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return []

def create_pdf_from_text(text):
    pdf = FPDF()
    pdf.add_page()
    try:
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.set_font("DejaVu", size=12)
    except RuntimeError:
        print("DejaVu font not found, falling back to Arial.")
        pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    return bytes(pdf.output(dest='S'))

def create_docx_from_text(text):
    doc = Document()
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

st.title("🗣️ Language Translation Agent")
st.markdown("**Translate Text or Documents using your chosen AI model.**")

st.markdown("Language and Model Selection")
col_lang, col_model, col_empty = st.columns([1, 1, 2])

languages = [
    "Arabic", "Bengali", "Chinese (Simplified)", "Chinese (Traditional)", 
    "Dutch", "English", "French", "German", "Hindi", "Indonesian", 
    "Italian", "Japanese", "Korean", "Malay", "Marathi", "Polish", 
    "Portuguese (Brazil)", "Punjabi", "Russian", "Spanish", "Tamil", 
    "Telugu", "Thai", "Turkish", "Urdu", "Vietnamese"
]
languages.sort()

try:
    default_lang_index = languages.index(st.session_state.target_language)
except ValueError:
    languages.append(st.session_state.target_language)
    languages.sort()
    default_lang_index = languages.index(st.session_state.target_language)

with col_lang:
    st.session_state.target_language = st.selectbox(
        "Select Target Language", languages, index=default_lang_index, label_visibility="collapsed"
    )

valid_model_names = []
if "custom_model" in st.secrets:
    valid_model_names.append("custom/gemma-3-4b-it") 
    
if st.session_state.gemini_api_configured:
    valid_model_names.extend([
        "models/gemini-2.5-pro",
        "models/gemini-2.5-flash"
    ])

with col_model:
    if not valid_model_names:
        st.error("No models available.")
        st.stop()
    try:
        default_model_index = valid_model_names.index(st.session_state.selected_model)
    except ValueError:
        default_model_index = 0

    st.session_state.selected_model = st.selectbox(
        "Select Model",
        options=valid_model_names,
        index=default_model_index,
        format_func=lambda name: name.replace("models/", "Gemini: ").replace("custom/", "Custom: "),
        label_visibility="collapsed"
    )

with st.expander("Advanced Options: Define Your Prompt Here"):
    st.text_area(
        "Edit the Prompt Instructions Below:",
        key="prompt_instructions",
        height=150,
        label_visibility="collapsed"
    )
st.divider()

st.subheader("Text Translation")
col1a, col1b = st.columns([1, 1], gap="large")

with col1a:
    input_text = st.text_area("Enter English Text Here:", height=300, key="text_input")
    if st.button("Translate Text", use_container_width=True, type="primary", disabled=not input_text):
        instructions = st.session_state.prompt_instructions.strip()
        spinner_model_name = st.session_state.selected_model.replace('models/', 'Gemini: ').replace("custom/", "Custom: ")
        
        with st.spinner(f"Translating using {spinner_model_name}..."):
            translated_output = translate_text(
                st.session_state.selected_model, 
                input_text, 
                st.session_state.target_language, 
                instructions
            )
            if translated_output is not None:
                st.session_state.text_translation_result = translated_output
                st.session_state.doc_translation_result = ""
                st.session_state.doc_info = {"name": None, "type": None}
                
                st.session_state.editing_translation = False
                st.session_state.current_source_text = input_text
                st.session_state.text_output_edit = translated_output
            else:
                st.error("Translation failed. Please check the console for errors.")

with col1b:
    if st.session_state.editing_translation:
        st.text_area(
            "Edit your translation:",
            height=300, 
            key="text_output_edit" 
        )
        if st.button("Save Correction", use_container_width=True, type="primary"):
            db_save_correction(
                db_pool,
                st.session_state.target_language,
                st.session_state.current_source_text,
                st.session_state.text_translation_result, 
                st.session_state.text_output_edit 
            )
            st.session_state.text_translation_result = st.session_state.text_output_edit
            st.session_state.editing_translation = False
            st.rerun()
            
    else:
        st.text_area(
            "Translated Text",
            value=st.session_state.text_translation_result or "Translation will appear here...",
            height=300, disabled=True, key="text_output"
        )
        if st.session_state.text_translation_result:
            if st.button("Edit Translation", use_container_width=True):
                st.session_state.editing_translation = True
                st.rerun()

st.divider()

st.subheader("Document Translation")
col2a, col2b = st.columns([1, 1], gap="large")

with col2a:
    uploaded_file = st.file_uploader("Upload a Document", type=["pdf", "docx", "txt"], key="file_uploader")
    if st.button("Translate Document", use_container_width=True, type="primary", disabled=uploaded_file is None):
        if uploaded_file is not None:
            instructions = st.session_state.prompt_instructions.strip()
            st.session_state.doc_info['name'] = uploaded_file.name
            file_ext = os.path.splitext(uploaded_file.name)[-1].lower().strip('.')
            st.session_state.doc_info['type'] = file_ext

            source_chunks = []
            if file_ext == "pdf": source_chunks = extract_text_from_pdf(uploaded_file)
            elif file_ext == "docx": source_chunks = extract_text_from_docx(uploaded_file)
            elif file_ext == "txt": source_chunks = extract_text_from_txt(uploaded_file)

            if source_chunks:
                translated_chunks = []
                total_chunks = math.ceil(len(source_chunks) / CHUNK_SIZE)
                progress_bar = st.progress(0, text=f"Translating chunk 1 of {total_chunks}...")

                for i in range(0, len(source_chunks), CHUNK_SIZE):
                    chunk_group = "\n".join(source_chunks[i:i + CHUNK_SIZE])
                    current_chunk_number = (i // CHUNK_SIZE) + 1
                    progress_text = f"Translating chunk {current_chunk_number} of {total_chunks}..."
                    progress_bar.progress(current_chunk_number / total_chunks, text=progress_text)
                    spinner_model_name = st.session_state.selected_model.replace('models/', 'Gemini: ').replace("custom/", "Custom: ")
                    
                    with st.spinner(f"Translating chunk {current_chunk_number}/{total_chunks} using {spinner_model_name}..."):
                        translated_output = translate_text(
                            st.session_state.selected_model, 
                            chunk_group, 
                            st.session_state.target_language, 
                            instructions
                        )
                    if translated_output is not None:
                        translated_chunks.append(translated_output)
                    else:
                        st.error(f"Failed to translate chunk {current_chunk_number}. Skipping.")
                        
                st.session_state.doc_translation_result = "\n\n".join(translated_chunks)
                st.session_state.text_translation_result = ""
                st.session_state.editing_translation = False 
                progress_bar.progress(1.0, text="Translation complete!")
            else:
                st.warning("Could not extract text from the document. The file might be empty or corrupted.")

with col2b:
    st.text_area(
        "Translated Document",
        value=st.session_state.doc_translation_result or "Document translation will appear here...",
        height=300,
        key="doc_output",
        disabled=True
    )
    if st.session_state.doc_translation_result:
        original_name = os.path.splitext(st.session_state.doc_info.get('name', 'document'))[0]
        file_ext = st.session_state.doc_info.get('type', 'txt')
        
        dl_cols = st.columns(3 if file_ext in ['pdf', 'docx'] else 1)
        
        with dl_cols[0]:
            st.download_button(
                label="Download as .txt",
                data=st.session_state.doc_translation_result.encode('utf-8'),
                file_name=f"translated_{original_name}.txt",
                mime="text/plain", 
                use_container_width=True
            )
        
        if file_ext == 'docx':
            with dl_cols[1]:
                st.download_button(
                    label="Download as .docx",
                    data=create_docx_from_text(st.session_state.doc_translation_result),
                    file_name=f"translated_{original_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        if file_ext == 'pdf':
            with dl_cols[1]:
                st.download_button(
                    label="Download as .pdf",
                    data=create_pdf_from_text(st.session_state.doc_translation_result),
                    file_name=f"translated_{original_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

